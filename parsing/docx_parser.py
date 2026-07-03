"""
.docx / .docm parser — single deterministic method: python-docx.

.docm is the same OOXML zip container as .docx plus a macro part python-docx
simply ignores; no special-casing needed (confirmed empirically against the
3 real .docm files in this corpus during implementation).
"""
from pathlib import Path

from docx import Document

from .common import ImageWriter, ParsedBlock, ParsedDocument

_DRAWING_TAGS = ("}drawing", "}pict")
_BLIP_TAG = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
_EMBED_ATTR = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"


_CORE_PROPS = (
    "author", "category", "comments", "content_status", "created", "identifier",
    "keywords", "language", "last_modified_by", "last_printed", "modified",
    "revision", "subject", "title", "version",
)


def parse(path: Path, image_writer: ImageWriter) -> ParsedDocument:
    docx = Document(str(path))
    parsed = ParsedDocument(source_file=path.name, doc_type=path.suffix.lstrip("."))
    cp = docx.core_properties
    parsed.metadata = {name: str(getattr(cp, name)) for name in _CORE_PROPS if getattr(cp, name, None)}

    image_rels = {}
    for rel in docx.part.rels.values():
        if "image" in rel.reltype:
            try:
                image_rels[rel.rId] = rel.target_part.blob
            except Exception:
                pass

    for para in docx.paragraphs:
        style = (para.style.name or "").lower() if para.style else ""
        text = para.text.strip()

        if style.startswith("heading"):
            try:
                level = int(style.replace("heading", "").strip() or "1")
            except ValueError:
                level = 2
            if text:
                parsed.blocks.append(ParsedBlock(kind="heading", content=text, level=level))
            continue

        for run in para.runs:
            for child in run._element.iter():
                if not any(child.tag.endswith(t) for t in _DRAWING_TAGS):
                    continue
                for blip in child.findall(f".//{_BLIP_TAG}"):
                    embed = blip.get(_EMBED_ATTR)
                    if embed and embed in image_rels:
                        image_id = image_writer.add(image_rels[embed], "png")
                        parsed.blocks.append(ParsedBlock(kind="image", image_id=image_id))

        if text:
            parsed.blocks.append(ParsedBlock(kind="text", content=text))

    for table in docx.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if rows:
            parsed.blocks.append(ParsedBlock(kind="table", content=rows))

    parsed.units_total = 1  # docx has no natural page count without rendering
    return parsed
